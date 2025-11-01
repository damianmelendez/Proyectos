from docx import Document
from docx.shared import Inches

# --------------------------------------------------------------------------------------
# FUNCIONES AUXILIARES: CONVERSION DE NUMEROS A LETRAS
# NOTA: Esta es una función simplificada para quetzales. Para un sistema de producción
# se recomienda usar librerías específicas que manejen correctamente grandes cantidades.
# --------------------------------------------------------------------------------------

def numero_a_letras(numero):
    """Convierte un número a su representación en Quetzales (simplificado)."""
    # Solo maneja números enteros y decimales comunes para este ejemplo
    try:
        entero = int(numero)
        decimal = int(round((numero - entero) * 100))
        
        # Mapa de números básicos (simplificado)
        unidades = ["", "UN", "DOS", "TRES", "CUATRO", "CINCO", "SEIS", "SIETE", "OCHO", "NUEVE"]
        decenas = ["", "DIEZ", "VEINTE", "TREINTA", "CUARENTA", "CINCUENTA", "SESENTA", "SETENTA", "OCHENTA", "NOVENTA"]
        
        # Una lógica muy simple para la parte entera (solo hasta 999)
        if entero < 10:
            parte_entera_letra = unidades[entero]
        elif entero < 100:
            parte_entera_letra = decenas[entero // 10]
            if entero % 10 != 0:
                parte_entera_letra += " Y " + unidades[entero % 10]
        else:
             # Para la renta y depósito, el usuario ingresará el valor en letras directamente
             # para evitar la complejidad de una librería completa.
             return input(f"Ingresa el monto de Q. {numero} en letras (ej: UN MIL TRESCIENTOS): ").upper()
             
        # Formato de la parte decimal
        if decimal == 0:
            return parte_entera_letra.upper() + " QUETZALES EXACTOS"
        else:
            return parte_entera_letra.upper() + f" QUETZALES CON {decimal}/100"
            
    except:
        # En caso de error o número complejo, pedimos el valor en letras al usuario.
        return input(f"Ingresa el monto de Q. {numero} en letras (ej: UN MIL TRESCIENTOS): ").upper()


# --------------------------------------------------------------------------------------
# PLANTILLA DEL CONTRATO (El texto que definimos en el punto 1)
# --------------------------------------------------------------------------------------

CONTRATO_PLANTILLA = """
POR EL PRESENTE DOCUMENTO: {{NOMBRE_ARRENDANTE_COMPLETO}}
de {{EDAD_ARRENDANTE}} años de edad, {{ESTADO_CIVIL_ARRENDANTE}}, de profesión {{PROFESION_ARRENDANTE}}, Guatemalteca.
Titular del Documento Personal de Identificación código único: {{DPI_ARRENDANTE_NUMERO}}, extendido por el Registro Nacional de las personas de la república de Guatemala quien de hoy en adelante se le denominará EL ARRENDANTE Y {{NOMBRE_ARRENDATARIO_COMPLETO}}, de {{EDAD_ARRENDATARIO}} años de edad, {{ESTADO_CIVIL_ARRENDATARIO}}, de profesión {{PROFESION_ARRENDATARIO}}, Guatemalteca, Titular del Documento Personal de Identificación código único: {{DPI_ARRENDATARIO_NUMERO}}, extendido por el Registro Nacional de las Personas de ta República de Guatemala.
Y quien de hoy en adelante se le denominará: EL ARRENDATARIO.
Aseguramos ser de los datos de identificación personal anteriormente indicados, hallarnos en el libre ejercicio de nuestros derechos civiles y que por este acto venimos a realizar el contrato de Arrendamiento de conformidad con las siguientes cláusulas: PRIMERA: {{NOMBRE_ARRENDANTE_COMPLETO}}, expreso que soy legítima propietaria el bien inmueble ubicado en {{DIRECCION_INMUEBLE}} y entrego en calidad de arrendamiento a: {{NOMBRE_ARRENDATARIO_COMPLETO}}, de acuerdo a las siguientes condiciones: a) PLAZO: El arrendamiento es de {{PLAZO_MESES_NUMERO}} ({{PLAZO_MESES_LETRA}}) a partir del día {{DIA_INICIO}} de {{MES_INICIO}} del año {{ANIO_INICIO}} y finaliza el {{DIA_FINAL}} de {{MES_FINAL}} del año del {{ANIO_FINAL}} pudiendo prorrogarse a voluntad de ambas partes siempre y cuando estén al día en sus pagos y exista un cruce de cartas con veinte días de anticipación y con firma de recepción y debe entenderse que si se recibiera la renta del mes siguiente al vencimiento del plazo será únicamente para el mes que se cancela y así sucesivamente, en el entendido de que
el presente contrato no se convertirá en indefinido y verbal b) RENTA: La renta será de {{RENTA_NUMERO_LETRAS}} (Q. {{RENTA_NUMERO}}) QUETZALES, se puede renovar el contrato con las estipulaciones anteriormente indicadas, los cuales deberán ser pagados en forma anticipada y sin necesidad de cobro ni requerimiento alguno en la residencia del arrendante ya conocida por el arrendatario, dentro de los primeros cinco días de cada mes.
Así mismo dentro del pago de la renta No se encuentra incluido Ningún Servicio y cualquier otro servicio corre a cuenta del Arrendatario c) DEPOSITO: Se deja un depósito de {{DEPOSITO_NUMERO_LETRAS}} (Q. {{DEPOSITO_NUMERO}}) QUETZALES EXACTOS.
Así mismo el Arrendatario al no cancelar la renta en el tiempo estipulado, tendrá una sanción económica de {{SANCION_DIARIA_LETRAS}} (Q. {{SANCION_DIARIA_NUMERO}}) por día atrasado d) DESTINO: El inmueble se destinará para {{DESTINO_INMUEBLE}} y no debe dársele otro destino sin el consentimiento expreso del arrendante, e) PROHIBIONES: queda prohibido, subarrendar el inmueble así como tener substancias prohibidas por la ley o personas que atenten contra la seguridad del estado o materiales inflamables, explosivas, salitrosas, corrosivas o que produzcan humedad o deterioren el inmueble, así como al Analizar el contrato no se da derecho de llave ni reclamar haber acreditado el local, f)
PAGO DE SERVICIOS: El local se entrega con todos sus servicios en buen estado y así deberá devolverse, únicamente con el deterioro natural.
El inquilino debe cancelar los servicios de energía eléctrica, agua potable, Extracción de basura y optaros servicios que solicite el arrendatario, g) REPARACIONES Y MEJORAS: El inquilino se compromete a entregar el inmueble en el mismo estado que lo recibe, limpio y con todos sus accesorios completos y cualquier reparación o mejora que se haga debe contar con la autorización del arrendante y de todas maneras las mejoras que se hagan queda a beneficio del inmueble h) RESCISIÓN: pl contrato se rescinde por 1) incumplimiento de parte del arrendatario a las cláusulas citadas anteriormente y dará lugar a la desocupación
inmediata del inmueble y al pago de los daños y perjuicios que se deriven de este contrato I) El arrendatario renuncia al fuero de su
domicilio y señala lugar para recibir citaciones y notificaciones el mismo local que recibe hoy en arrendamiento de lo contrario se tendrán como bien hechas las citaciones y notificaciones y emplazamientos que se hagan, en caso no diere aviso del cambio de dirección.
J) OTROS: desde ya el arrendatario acepta como bien hechas y válidas las cuentas que se le presenten y en el caso de solicitar la intervención del órgano jurisdiccional debe de cancelar los gastos que se ocasiones K) El inmueble se encuentra en buen estado de habitabilidad y con todos sus servicios en perfecto estado y así debe devolverse.
SEGUNDA? Los comparecientes expresamos que en los términos relacionados aceptamos el presente contrato de arrendamiento y enterados de su contenido, objeto, validez y demás efectos legales, lo firmamos en la {{CIUDAD_FIRMA}} el {{FECHA_FIRMA_ESCRITURA}}.

***

En la {{CIUDAD_FIRMA}}, el {{FECHA_LEGALIZACION_FIRMA}}, Yo el Infrascrito Notario doy fe que las firmas que anteceden son AUTENTICAS por haber sido puestas el día de hoy en mi presencia por: {{NOMBRE_ARRENDANTE_COMPLETO}}, Titular del Documento Personal de Identificación código único: {{DPI_ARRENDANTE_NUMERO}}, extendido por el Registro Nacional de las personas de la república de Guatemala y {{NOMBRE_ARRENDATARIO_COMPLETO}}, Titular del Documento Personal de Identificación código único: {{DPI_ARRENDATARIO_NUMERO}}, extendido por el Registro Nacional de las Personas de la República de Guatemala y quienes previa lectura de la presente acta de legalización de firma la ratifican.
Aceptan y vuelven a firmar.
"""


def solicitar_datos():
    """Solicita al usuario todos los datos necesarios para rellenar el contrato."""
    print("--- 🏠 INGRESO DE DATOS PARA EL CONTRATO DE ARRENDAMIENTO ---")
    
    # Datos de Arrendante (Propietario)
    print("\n[ DATOS DEL ARRENDANTE ]")
    datos = {}
    datos['NOMBRE_ARRENDANTE_COMPLETO'] = input("Nombre completo del Arrendante: ")
    datos['EDAD_ARRENDANTE'] = input("Edad del Arrendante (en números): ")
    datos['ESTADO_CIVIL_ARRENDANTE'] = input("Estado civil del Arrendante (ej: casada): ")
    datos['PROFESION_ARRENDANTE'] = input("Profesión del Arrendante: ")
    datos['DPI_ARRENDANTE_NUMERO'] = input("DPI del Arrendante (con espacios o guiones, ej: 2329 78204 0301): ")

    # Datos de Arrendatario (Inquilino)
    print("\n[ DATOS DEL ARRENDATARIO ]")
    datos['NOMBRE_ARRENDATARIO_COMPLETO'] = input("Nombre completo del Arrendatario: ")
    datos['EDAD_ARRENDATARIO'] = input("Edad del Arrendatario (en números): ")
    datos['ESTADO_CIVIL_ARRENDATARIO'] = input("Estado civil del Arrendatario (ej: casada): ")
    datos['PROFESION_ARRENDATARIO'] = input("Profesión del Arrendatario: ")
    datos['DPI_ARRENDATARIO_NUMERO'] = input("DPI del Arrendatario (con espacios o guiones, ej: 2347 75866 0101): ")

    # Datos del Inmueble y Plazos
    print("\n[ DATOS DEL INMUEBLE Y CONDICIONES ]")
    datos['DIRECCION_INMUEBLE'] = input("Dirección completa del inmueble: ")
    datos['DESTINO_INMUEBLE'] = input("Destino del inmueble (ej: OFICINA, VIVIENDA): ").upper()
    
    # Plazo
    datos['PLAZO_MESES_NUMERO'] = input("Plazo del arrendamiento en números (ej: 12): ")
    datos['PLAZO_MESES_LETRA'] = input("Plazo del arrendamiento en letras (ej: DOCE MESES): ").upper()
    
    # Fechas de inicio y fin (Recuerda ingresar solo las palabras o números que cambian)
    print("\n[ FECHAS DEL CONTRATO ]")
    datos['DIA_INICIO'] = input("Día de inicio del contrato (ej: UNO): ").upper()
    datos['MES_INICIO'] = input("Mes de inicio del contrato (ej: Abril): ").capitalize()
    datos['ANIO_INICIO'] = input("Año de inicio del contrato (ej: dos mil veinticuatro): ")
    
    datos['DIA_FINAL'] = input("Día de finalización del contrato (ej: treinta y uno): ").upper()
    datos['MES_FINAL'] = input("Mes de finalización del contrato (ej: Mayo): ").capitalize()
    datos['ANIO_FINAL'] = input("Año de finalización del contrato (ej: dos mil veinticuatro): ")
    
    # Renta, Depósito y Sanción
    print("\n[ MONTOS DEL CONTRATO (QUETZALES) ]")
    datos['RENTA_NUMERO'] = input("Monto de la Renta (en números, ej: 1300): ")
    datos['RENTA_NUMERO_LETRAS'] = input("Monto de la Renta (en letras, ej: UN MIL TRESCIENTOS): ").upper()
    
    datos['DEPOSITO_NUMERO'] = input("Monto del Depósito (en números, ej: 1300): ")
    datos['DEPOSITO_NUMERO_LETRAS'] = input("Monto del Depósito (en letras, ej: UN MIL TRESCIENTOS): ").upper()
    
    datos['SANCION_DIARIA_NUMERO'] = input("Monto de la Sanción Diaria (en números, ej: 30): ")
    datos['SANCION_DIARIA_LETRAS'] = input("Monto de la Sanción Diaria (en letras, ej: TREINTA): ").upper()
    
    # Fechas de Firma y Legalización
    print("\n[ FECHAS DE ESCRITURA Y LEGALIZACIÓN ]")
    datos['CIUDAD_FIRMA'] = input("Ciudad de la firma (ej: Ciudad de la Antigua Guatemala del Departamento de Sacatepéquez): ")
    datos['FECHA_FIRMA_ESCRITURA'] = input("Fecha de firma de la Escritura (ej: uno de Abril del año dos mil veinticuatro): ")
    datos['FECHA_LEGALIZACION_FIRMA'] = input("Fecha de legalización de firma (ej: veinticinco de enero del año dos mil veinticuatro): ")

    return datos

def generar_contrato_docx(datos):
    """Genera el documento DOCX con los datos ingresados."""
    
    # 1. Crear el texto final reemplazando los marcadores
    texto_final = CONTRATO_PLANTILLA
    for marcador, valor in datos.items():
        # Reemplazar todos los marcadores con el valor correspondiente
        texto_final = texto_final.replace("{{" + marcador + "}}", valor)
    
    # 2. Crear un nuevo documento Word
    documento = Document()
    
    # Dividir el texto en párrafos usando el salto de línea
    parrafos = texto_final.split('\n')
    
    for parrafo_texto in parrafos:
        # Usar strip() para eliminar espacios en blanco al inicio/final del párrafo
        if parrafo_texto.strip() == '***':
            # Insertar un separador visual o un salto de línea si es necesario
            documento.add_paragraph()
            documento.add_paragraph('************************************************************************************************************************')
            documento.add_paragraph()
            continue

        # Evitar párrafos vacíos si hay múltiples saltos de línea
        if parrafo_texto.strip():
            # Añadir el texto al documento. Aquí puedes añadir lógica para formato (negritas, etc.)
            documento.add_paragraph(parrafo_texto)

    # 3. Guardar el documento
    nombre_archivo = f"Contrato_{datos['NOMBRE_ARRENDATARIO_COMPLETO'].replace(' ', '_')}_{datos['ANIO_INICIO']}.docx"
    documento.save(nombre_archivo)
    print(f"\n🎉 ¡Éxito! El contrato ha sido generado y guardado como: {nombre_archivo}")
    print("Recuerda que deberás abrir el archivo para añadir las firmas escaneadas si lo necesitas.")

if __name__ == "__main__":
    datos_contrato = solicitar_datos()
    generar_contrato_docx(datos_contrato)