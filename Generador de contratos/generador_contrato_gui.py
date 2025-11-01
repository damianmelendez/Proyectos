import tkinter as tk
from tkinter import ttk, messagebox
from docx import Document
import sys

# --------------------------------------------------------------------------------------
# LÓGICA DE GENERACIÓN DE CONTRATO (Mantenemos la plantilla y la función de reemplazo)
# --------------------------------------------------------------------------------------

CONTRATO_PLANTILLA = """
POR EL PRESENTE DOCUMENTO: {{NOMBRE_ARRENDANTE_COMPLETO}}
de {{EDAD_ARRENDANTE}} años de edad, {{ESTADO_CIVIL_ARRENDANTE}}, de profesión {{PROFESION_ARRENDANTE}}, Guatemalteca.
Titular del Documento Personal de Identificación código único: {{DPI_ARRENDANTE_NUMERO}}, extendido por el Registro Nacional de las personas de la república de Guatemala quien de hoy en adelante se le denominará EL ARRENDANTE Y {{NOMBRE_ARRENDATARIO_COMPLETO}}, de {{EDAD_ARRENDATARIO}} años de edad, {{ESTADO_CIVIL_ARRENDATARIO}}, de profesión {{PROFESION_ARRENDATARIO}}, Guatemalteca, Titular del Documento Personal de Identificación código único: {{DPI_ARRENDATARIO_NUMERO}}, extendido por el Registro Nacional de las Personas de ta República de Guatemala.
Y quien de hoy en adelante se le denominará: EL ARRENDATARIO.
Aseguramos ser de los datos de identificación personal anteriormente indicados, hallarnos en el libre ejercicio de nuestros derechos civiles y que por este acto venimos a realizar el contrato de Arrendamiento de conformidad con las siguientes cláusulas: PRIMERA: {{NOMBRE_ARRENDANTE_COMPLETO}}, expreso que soy legítima propietaria el bien inmueble ubicado en {{DIRECCION_INMUEBLE}} y entrego en calidad de arrendamiento a: {{NOMBRE_ARRENDATARIO_COMPLETO}}, de acuerdo a las siguientes condiciones: a) PLAZO: El arrendamiento es de {{PLAZO_MESES_NUMERO}} ({{PLAZO_MESES_LETRA}}) a partir del día {{DIA_INICIO}} de {{MES_INICIO}} del año {{ANIO_INICIO}} y finaliza el {{DIA_FINAL}} de {{MES_FINAL}} del año del {{ANIO_FINAL}} pudiendo prorrogarse a voluntad de ambas partes siempre y cuando estén al día en sus pagos y exista un cruce de cartas con veinte días de anticipación y con firma de recepción y debe entenderse que si se recibiera la renta del mes siguiente al vencimiento del plazo será únicamente para el mes que se cancela y así sucesivamente, en el entendido de que
el presente contrato no se convertirá en indefinido y verbal b) RENTA: La renta será de {{RENTA_NUMERO_LETRAS}} (Q. {{RENTA_NUMERO}}) QUETZALES, se puede renovar el contrato con las estipulaciones anteriormente indicadas, los cuales deberán ser pagados en forma anticipada y sin necesidad de cobro ni requerimiento alguno en la residencia del arrendante ya conocida por el arrendatario, dentro de los primeros cinco días de cada mes.
Así mismo dentro del pago de la renta No se encuentra incluido Ningún Servicio y cualquier otro servicio corre a cuenta del Arrendatario c) DEPOSITO: Se deja un depósito de {{DEPOSITO_NUMERO_LETRAS}} (Q. {{DEPOSITO_NUMERO}}) QUETZALES EXACTOS.
Así mismo el Arrendatario al no cancelar la renta en el tiempo estipulado, tendrá una sanción económica de {{SANCION_DIARIA_LETRAS}} (Q. {{SANCION_DIARIA_NUMERO}}) por día atrasado d) DESTINO: El inmueble se destinará para {{DESTINO_INMUEBLE}} y no debe dársele otro destino sin el consentimiento expreso del arrendante, e) PROHIBIONES: queda prohibido, subarrendar el inmueble así como tener substancias prohibidas por la ley o personas que atenten contra la seguridad del estado o materiales inflamables, explosivas, salitrosas, corrosivas o que produzcan humedad o deterioren el inmueble, así como al Analizar el contrato no se da derecho de llave ni reclamar haber acreditado el local, f)
PAGO DE SERVICIOS: El local se entrega con todos sus servicios en buen estado y así deberá devolverse, únicamente con el deterioro natural.
El inquilino debe cancelar los servicios de energía eléctrica, agua potable, Extracción de basura y optaros servicios que solicite el arrendatario, g) REPARACIONES Y MEJORAS: El inquilino se compromete a entregar el inmueble en el mismo estado que lo recibe, limpio y con todos sus accesorios completos y cualquier reparación o mejora que se haga debe contar con la autorización del arrendante y de todas maneras las mejoras que se hagan queda a beneficio del inmueble h) RESCISIÓN: pl contrato se rescinde por 1) incumplimiento de parte del arrendatario a las cláusulas citadas anteriormente y dará lugar a la desocupación
inmediata del inmueble y al pago de los daños y perjuicios que se deriven de este contrato I) El arrendatario renuncia al fuero de su
domicilio y señala lugar para recibir citaciones y notificaciones el mismo local que recibe hoy en arrendamiento de lo contrario se tendrán como bien hechas las citaciones y notificaciones y emplazamientos que se hagan, en caso no diere aviso del cambio de dirección.
J) OTROS: desde ya el arrendatario acepta como bien hechas y válidas las cuentas que se le presenten y en el caso de solicitar la intervención del órgano jurisdiccional debe de cancelar los gastos que se ocasiones K) El inmueble se encuentra en buen estado de habitabilidad y con todos sus servicios en perfecto estado y así debe devolverse.
SEGUNDA? Los comparecientes expresamos que en los términos relacionados aceptamos el presente contrato de arrendamiento y enterados de su contenido, objeto, validez y demás efectos legales, lo firmamos en la {{CIUDAD_FIRMA}} el {{FECHA_FIRMA_ESCRITURA}}.

***

En la {{CIUDAD_FIRMA}}, el {{FECHA_LEGALIZACION_FIRMA}}, Yo el Infrascrito Notario doy fe que las firmas que anteceden son AUTENTICAS por haber sido puestas el día de hoy en mi presencia por: {{NOMBRE_ARRENDANTE_COMPLETO}}, Titular del Documento Personal de Identificación código único: {{DPI_ARRENDANTE_NUMERO}}, extendido por el Registro Nacional de las personas de la república de Guatemala y {{NOMBRE_ARRENDATARIO_COMPLETO}}, Titular del Documento Personal de Identificación código único: {{DPI_ARRENDATARIO_NUMERO}}, extendido por el Registro Nacional de las Personas de la República de Guatemala y quienes previa lectura de la presente acta de legalización de firma la ratifican.
Aceptan y vuelven a firmar.
"""


def generar_contrato_docx(datos):
    """Genera el documento DOCX con los datos ingresados."""
    
    # 1. Crear el texto final reemplazando los marcadores
    texto_final = CONTRATO_PLANTILLA
    for marcador, valor in datos.items():
        # Reemplazar todos los marcadores con el valor correspondiente
        texto_final = texto_final.replace("{{" + marcador + "}}", str(valor))
    
    # 2. Crear un nuevo documento Word
    documento = Document()
    
    # Dividir el texto en párrafos usando el salto de línea
    parrafos = texto_final.split('\n')
    
    for parrafo_texto in parrafos:
        if parrafo_texto.strip() == '***':
            # Separador visual
            documento.add_paragraph()
            documento.add_paragraph('*** PARTE DE LEGALIZACIÓN DE FIRMA ***')
            documento.add_paragraph()
            continue

        if parrafo_texto.strip():
            documento.add_paragraph(parrafo_texto)

    # 3. Guardar el documento
    # Usamos el nombre del Arrendatario y el Año para nombrar el archivo
    nombre_archivo = f"Contrato_{datos['NOMBRE_ARRENDATARIO_COMPLETO'].replace(' ', '_').replace('.', '')}_{datos['ANIO_INICIO'].replace(' ', '_')}.docx"
    try:
        documento.save(nombre_archivo)
        messagebox.showinfo("Éxito", f"¡Contrato generado con éxito!\nGuardado como: {nombre_archivo}")
    except Exception as e:
        messagebox.showerror("Error al guardar", f"No se pudo guardar el archivo DOCX: {e}")


# --------------------------------------------------------------------------------------
# INTERFAZ GRÁFICA (TKINTER)
# --------------------------------------------------------------------------------------

class ContratoApp:
    def __init__(self, master):
        self.master = master
        master.title("Generador Automático de Contratos")
        
        # Diccionario para almacenar las entradas de texto
        self.entries = {}
        
        # Definición de los campos agrupados para la GUI
        self.campos = {
            "Datos del Arrendante (Propietario)": [
                ('Nombre completo del Arrendante', 'NOMBRE_ARRENDANTE_COMPLETO'),
                ('Edad (ej: 69)', 'EDAD_ARRENDANTE'),
                ('Estado Civil (ej: casada)', 'ESTADO_CIVIL_ARRENDANTE'),
                ('Profesión', 'PROFESION_ARRENDANTE'),
                ('DPI (ej: 2329 78204 0301)', 'DPI_ARRENDANTE_NUMERO'),
            ],
            "Datos del Arrendatario (Inquilino)": [
                ('Nombre completo del Arrendatario', 'NOMBRE_ARRENDATARIO_COMPLETO'),
                ('Edad (ej: 46)', 'EDAD_ARRENDATARIO'),
                ('Estado Civil (ej: casada)', 'ESTADO_CIVIL_ARRENDATARIO'),
                ('Profesión', 'PROFESION_ARRENDATARIO'),
                ('DPI (ej: 2347 75866 0101)', 'DPI_ARRENDATARIO_NUMERO'),
            ],
            "Datos del Inmueble y Plazos": [
                ('Dirección del Inmueble', 'DIRECCION_INMUEBLE'),
                ('Destino (ej: OFICINA)', 'DESTINO_INMUEBLE'),
                ('Plazo en meses (Número, ej: 2)', 'PLAZO_MESES_NUMERO'),
                ('Plazo en meses (Letras, ej: DOS MESES)', 'PLAZO_MESES_LETRA'),
            ],
            "Fechas del Plazo": [
                ('Día de Inicio (Letras, ej: UNO)', 'DIA_INICIO'),
                ('Mes de Inicio (ej: Abril)', 'MES_INICIO'),
                ('Año de Inicio (Letras, ej: dos mil veinticuatro)', 'ANIO_INICIO'),
                ('Día Final (Letras, ej: TREINTA Y UNO)', 'DIA_FINAL'),
                ('Mes Final (ej: Mayo)', 'MES_FINAL'),
                ('Año Final (Letras, ej: dos mil veinticuatro)', 'ANIO_FINAL'),
            ],
            "Montos (Quetzales)": [
                ('Renta (Número, ej: 1300)', 'RENTA_NUMERO'),
                ('Renta (Letras, ej: UN MIL TRESCIENTOS)', 'RENTA_NUMERO_LETRAS'),
                ('Depósito (Número)', 'DEPOSITO_NUMERO'),
                ('Depósito (Letras)', 'DEPOSITO_NUMERO_LETRAS'),
                ('Sanción Diaria (Número, ej: 30)', 'SANCION_DIARIA_NUMERO'),
                ('Sanción Diaria (Letras, ej: TREINTA)', 'SANCION_DIARIA_LETRAS'),
            ],
            "Fechas y Lugar de Firma": [
                ('Ciudad de Firma (Completo)', 'CIUDAD_FIRMA'),
                ('Fecha Firma Escritura (Letras, ej: uno de Abril del año...)', 'FECHA_FIRMA_ESCRITURA'),
                ('Fecha Legalización Notario (Letras, ej: veinticinco de enero del año...)', 'FECHA_LEGALIZACION_FIRMA'),
            ]
        }
        
        self.crear_widgets()

    def crear_widgets(self):
        # Crear un Canvas y un Scrollbar para manejar el desplazamiento
        canvas = tk.Canvas(self.master)
        scrollbar = ttk.Scrollbar(self.master, orient="vertical", command=canvas.yview)
        scrollable_frame = ttk.Frame(canvas)

        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(
                scrollregion=canvas.bbox("all")
            )
        )

        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)

        # Empaquetar Canvas y Scrollbar
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        
        # Llenar la ventana con los campos de entrada
        fila = 0
        for seccion, campos in self.campos.items():
            # Título de la sección
            ttk.Label(scrollable_frame, text=seccion, font=('Arial', 10, 'bold'), foreground='blue').grid(row=fila, column=0, columnspan=2, pady=(10, 5), sticky='w')
            fila += 1
            
            for etiqueta, clave in campos:
                # Etiqueta (Label)
                ttk.Label(scrollable_frame, text=f"{etiqueta}:").grid(row=fila, column=0, padx=5, pady=2, sticky='w')
                
                # Campo de Entrada (Entry)
                entrada = ttk.Entry(scrollable_frame, width=60)
                entrada.grid(row=fila, column=1, padx=5, pady=2, sticky='ew')
                self.entries[clave] = entrada # Guardar referencia para obtener el valor después
                fila += 1
        
        # Botón de Generar
        ttk.Button(scrollable_frame, text="Generar Contrato DOCX", command=self.obtener_y_generar).grid(row=fila, column=0, columnspan=2, pady=20)
        
        # Asegurar que las columnas se expandan correctamente
        scrollable_frame.grid_columnconfigure(1, weight=1)

    def obtener_y_generar(self):
        """Función que obtiene los datos de la GUI y llama al generador DOCX."""
        datos = {}
        # Obtener los valores de todas las entradas
        for clave, entrada in self.entries.items():
            valor = entrada.get().strip()
            if not valor:
                messagebox.showwarning("Faltan datos", f"El campo '{clave}' no puede estar vacío. Por favor, rellena toda la información.")
                return # Detiene la ejecución si falta algún campo

            # Los campos que deben ir en mayúsculas (ej: DESTINO, letras de plazo/montos) se convierten
            if any(k in clave for k in ['LETRA', 'DIA', 'DESTINO', 'ARRENDANTE', 'ARRENDATARIO', 'ESTADO_CIVIL']):
                datos[clave] = valor.upper()
            elif 'MES' in clave: # Meses solo la primera en mayúscula
                datos[clave] = valor.capitalize()
            else:
                datos[clave] = valor

        # Llamar a la función que genera el documento
        generar_contrato_docx(datos)

if __name__ == "__main__":
    # Asegurarse de que el color de fondo de la ventana sea claro para la apariencia de Tkinter
    root = tk.Tk()
    # Inicializar y correr la aplicación
    app = ContratoApp(root)
    root.mainloop()