import tkinter as tk
from tkinter import ttk, messagebox, simpledialog
from docx import Document
from docx.shared import Inches

# --------------------------------------------------------------------------------------
# 1. PLANTILLA DEL CONTRATO (Mantenida igual)
# --------------------------------------------------------------------------------------

CONTRATO_PLANTILLA = """
POR EL PRESENTE DOCUMENTO: {{NOMBRE_ARRENDANTE_COMPLETO}}
de {{EDAD_ARRENDANTE}} años de edad, {{ESTADO_CIVIL_ARRENDANTE}}, de profesión {{PROFESION_ARRENDANTE}}, Guatemalteca.
Titular del Documento Personal de Identificación código único: {{DPI_ARRENDANTE_NUMERO}}, extendido por el Registro Nacional de las personas de la república de Guatemala quien de hoy en adelante se le denominará EL ARRENDANTE Y {{NOMBRE_ARRENDATARIO_COMPLETO}}, de {{EDAD_ARRENDATARIO}} años de edad, {{ESTADO_CIVIL_ARRENDATARIO}}, de profesión {{PROFESION_ARRENDATARIO}}, Guatemalteca, Titular del Documento Personal de Identificación código único: {{DPI_ARRENDATARIO_NUMERO}}, extendido por el Registro Nacional de las Personas de ta República de Guatemala.
Y quien de hoy en adelante se le denominará: EL ARRENDATARIO.
Aseguramos ser de los datos de identificación personal anteriormente indicados, hallarnos en el libre ejercicio de nuestros derechos civiles y que por este acto venimos a realizar el contrato de Arrendamiento de conformidad con las siguientes cláusulas: PRIMERA: {{NOMBRE_ARRENDANTE_COMPLETO}}, expreso que soy legítima propietaria el bien inmueble ubicado en {{DIRECCION_INMUEBLE}} y entrego en calidad de arrendamiento a: {{NOMBRE_ARRENDATARIO_COMPLETO}}, de acuerdo a las siguientes condiciones: a) PLAZO: El arrendamiento es de {{PLAZO_MESES_NUMERO}} ({{PLAZO_MESES_LETRA}}) a partir del día {{DIA_INICIO}} de {{MES_INICIO}} del año {{ANIO_INICIO}} y finaliza el {{DIA_FINAL}} de {{MES_FINAL}} del año del {{ANIO_FINAL}} pudiendo prorrogarse a voluntad de ambas partes siempre y cuando estén al día en sus pagos y exista un cruce de cartas con veinte días de anticipación y con firma de recepción y debe entenderse que si se recibiera la renta del mes siguiente al vencimiento del plazo será únicamente para el mes que se cancela y así sucesivamente, en el entendido de que
el presente contrato no se convertirá en indefinido y verbal b) RENTA: La renta será de {{RENTA_NUMERO_LETRAS}} (Q. {{RENTA_NUMERO}}) QUETZALES, se puede renovar el contrato con las estipulaciones anteriormente indicadas, los cuales deberán ser pagados en forma anticipada y sin necesidad de cobro ni requerimiento alguno en la residencia del arrendante ya conocida por el arrendatario, dentro de los primeros cinco días de cada mes.
Así mismo dentro del pago de la renta No se encuentra incluido Ningún Servicio y cualquier otro servicio corre a cuenta del Arrendatario c) DEPOSITO: Se deja un depósito de {{DEPOSITO_NUMERO_LETRAS}} (Q. {{DEPOSITO_NUMERO}}) QUETZALES EXACTOS.
Así mismo el Arrendatario al no cancelar la renta en el tiempo estipulado, tendrá una sanción económica de {{SANCION_DIARIA_LETRAS}} (Q. {{SANCION_DIARIA_NUMERO}}) por día atrasado d) DESTINO: El inmueble se destinará para {{DESTINO_INMUEBLE}} y no debe dársele otro destino sin el consentimiento expreso del arrendante, e) PROHIBIONES: queda prohibido, subarrendar el inmueble así como tener substancias prohibidas por la ley o personas que atenten contra la seguridad del estado o materiales inflamables, explosivas, salitrosas, corrosivas o que produzcan humedad o deterioren el inmueble, así como al Analizar el contrato no se da derecho de llave ni reclamar haber acreditado el local, f)
PAGO DE SERVICIOS: El local se entrega con todos sus servicios en buen estado y así deberá devolverse, únicamente con el deterioro natural.
El inquilino debe cancelar los servicios de energía eléctrica, agua potable, Extracción de basura y optaros servicios que solicite el arrendatario, g) REPARACIONES Y MEJORAS: El inquilino se compromete a entregar el inmueble en el mismo estado que lo recibe, limpio y con todos sus accesorios completos y cualquier reparación o mejora que se haga debe contar con la autorización del arrendante y de todas maneras las mejoras que se hagan queda a beneficio del inmueble h) RESCISIÓN: pl contrato se rescinde por 1) incumplimiento de parte del arrendatario a las cláusulas citadas anteriormente y dará lugar a la desocupación
inmediata del inmueble y al pago de los daños y perjuicios que se deriven de este contrato I) El arrendatario renuncia al fuero de su
domicilio y señala lugar para recibir citaciones y notificaciones el mismo local que recibe hoy en arrendamiento de lo contrario se tendrán como bien hechas las citaciones y notificaciones y emplazamientos que se hagan, en caso no diere aviso del cambio de dirección.
J) OTROS: desde ya el arrendatario acepta como bien hechas y válidas las cuentas que se le presenten y en el caso de solicitar la intervención del órgano jurisdiccional debe de cancelar los gastos que se ocasiones K) El inmueble se encuentra en buen estado de habitabilidad y con todos sus servicios en perfecto estado y así debe devolverse.
SEGUNDA? Los comparecientes expresamos que en los términos relacionados aceptamos el presente contrato de arrendamiento y enterados de su contenido, objeto, validez y demás efectos legales, lo firmamos en la {{CIUDAD_FIRMA}} el {{FECHA_FIRMA_ESCRITURA}}.

***

En la {{CIUDAD_FIRMA}}, el {{FECHA_LEGALIZACION_FIRMA}}, Yo el Infrascrito Notario doy fe que las firmas que anteceden son AUTENTICAS por haber sido puestas el día de hoy en mi presencia por: {{NOMBRE_ARRENDANTE_COMPLETO}}, Titular del Documento Personal de Identificación código único: {{DPI_ARRENDANTE_NUMERO}}, extendido por el Registro Nacional de las personas de la república de Guatemala y {{NOMBRE_ARRENDATARIO_COMPLETO}}, Titular del Documento Personal de Identificación código único: {{DPI_ARRENDATARIO_NUMERO}}, extendido por el Registro Nacional de las Personas de la República de Guatemala y quienes previa lectura de la presente acta de legalización de firma la ratifican.
Aceptan y vuelven a firmar.
"""

# --------------------------------------------------------------------------------------
# 2. DEFINICIÓN DE CAMPOS PARA LA GUI
# --------------------------------------------------------------------------------------

CAMPOS = [
    # Sección Arrendante
    ("NOMBRE_ARRENDANTE_COMPLETO", "Nombre Completo Arrendante", "ARRENDANTE"),
    ("EDAD_ARRENDANTE", "Edad Arrendante (Números)", "ARRENDANTE"),
    ("ESTADO_CIVIL_ARRENDANTE", "Estado Civil Arrendante", "ARRENDANTE"),
    ("PROFESION_ARRENDANTE", "Profesión Arrendante", "ARRENDANTE"),
    ("DPI_ARRENDANTE_NUMERO", "DPI Arrendante (CUC)", "ARRENDANTE"),
    
    # Sección Arrendatario
    ("NOMBRE_ARRENDATARIO_COMPLETO", "Nombre Completo Arrendatario", "ARRENDATARIO"),
    ("EDAD_ARRENDATARIO", "Edad Arrendatario (Números)", "ARRENDATARIO"),
    ("ESTADO_CIVIL_ARRENDATARIO", "Estado Civil Arrendatario", "ARRENDATARIO"),
    ("PROFESION_ARRENDATARIO", "Profesión Arrendatario", "ARRENDATARIO"),
    ("DPI_ARRENDATARIO_NUMERO", "DPI Arrendatario (CUC)", "ARRENDATARIO"),
    
    # Sección Inmueble y Plazos
    ("DIRECCION_INMUEBLE", "Dirección Inmueble", "INMUEBLE/PLAZO"),
    ("DESTINO_INMUEBLE", "Destino Inmueble (MAYÚS)", "INMUEBLE/PLAZO"),
    ("PLAZO_MESES_NUMERO", "Plazo (Meses en Números)", "INMUEBLE/PLAZO"),
    ("PLAZO_MESES_LETRA", "Plazo (Meses en Letras)", "INMUEBLE/PLAZO"),
    
    # Sección Fechas de Plazo
    ("DIA_INICIO", "Día Inicio (LETRAS)", "FECHAS PLAZO"),
    ("MES_INICIO", "Mes Inicio", "FECHAS PLAZO"),
    ("ANIO_INICIO", "Año Inicio", "FECHAS PLAZO"),
    ("DIA_FINAL", "Día Fin (LETRAS)", "FECHAS PLAZO"),
    ("MES_FINAL", "Mes Fin", "FECHAS PLAZO"),
    ("ANIO_FINAL", "Año Fin", "FECHAS PLAZO"),
    
    # Sección Montos
    ("RENTA_NUMERO", "Renta (Números)", "MONTOS (Q)"),
    ("RENTA_NUMERO_LETRAS", "Renta (Letras MAYÚS)", "MONTOS (Q)"),
    ("DEPOSITO_NUMERO", "Depósito (Números)", "MONTOS (Q)"),
    ("DEPOSITO_NUMERO_LETRAS", "Depósito (Letras MAYÚS)", "MONTOS (Q)"),
    ("SANCION_DIARIA_NUMERO", "Sanción Diaria (Números)", "MONTOS (Q)"),
    ("SANCION_DIARIA_LETRAS", "Sanción Diaria (Letras MAYÚS)", "MONTOS (Q)"),
    
    # Sección Firma
    ("CIUDAD_FIRMA", "Ciudad de la Firma", "FECHAS FIRMA"),
    ("FECHA_FIRMA_ESCRITURA", "Fecha Escritura (LETRAS)", "FECHAS FIRMA"),
    ("FECHA_LEGALIZACION_FIRMA", "Fecha Legalización (LETRAS)", "FECHAS FIRMA"),
]

# --------------------------------------------------------------------------------------
# 3. LÓGICA DE GENERACIÓN DEL DOCX (Mantenida y corregida para negritas)
# --------------------------------------------------------------------------------------

def generar_contrato_docx(datos, filename="contrato_generado.docx"):
    """Genera el documento DOCX con los datos ingresados, aplicando formato de negritas."""
    
    try:
        documento = Document()
        
        # 1. Crear el texto final reemplazando los marcadores
        texto_final_plano = CONTRATO_PLANTILLA
        for marcador, valor in datos.items():
            texto_final_plano = texto_final_plano.replace("{{" + marcador + "}}", valor)
        
        # 2. Definir las frases clave en negrita
        negritas_fijas = [
            "POR EL PRESENTE DOCUMENTO:", "EL ARRENDANTE", "EL ARRENDATARIO", "PRIMERA:", 
            "a) PLAZO:", "b) RENTA:", "c) DEPOSITO:", "d) DESTINO:", "e) PROHIBIONES:", 
            "f) PAGO DE SERVICIOS:", "g) REPARACIONES Y MEJORAS:", "h) RESCISIÓN:", 
            "I)", "J) OTROS:", "K)", "SEGUNDA?", "AUTENTICAS"
        ]
        
        # Valores dinámicos que deben ir en negrita
        negritas_dinamicas = [
            datos['NOMBRE_ARRENDANTE_COMPLETO'],
            datos['NOMBRE_ARRENDATARIO_COMPLETO'],
            datos['PLAZO_MESES_NUMERO'],
            datos['PLAZO_MESES_LETRA'],
            datos['DIA_INICIO'],
            datos['MES_INICIO'],
            datos['ANIO_INICIO'],
            datos['DIA_FINAL'],
            datos['MES_FINAL'],
            datos['ANIO_FINAL'],
            datos['RENTA_NUMERO_LETRAS'], 
            f"(Q. {datos['RENTA_NUMERO']})",
            datos['DEPOSITO_NUMERO_LETRAS'],
            f"(Q. {datos['DEPOSITO_NUMERO']})",
            datos['SANCION_DIARIA_LETRAS'],
            f"(Q. {datos['SANCION_DIARIA_NUMERO']})",
            datos['CIUDAD_FIRMA'],
            datos['FECHA_FIRMA_ESCRITURA'],
            datos['FECHA_LEGALIZACION_FIRMA']
        ]
        
        palabras_en_negrita = list(set(negritas_fijas + negritas_dinamicas))
        palabras_en_negrita_ordenadas = sorted(palabras_en_negrita, key=len, reverse=True)


        # 3. Procesar cada párrafo (línea) para aplicar negritas
        parrafos_texto = texto_final_plano.split('\n')
        
        for parrafo_texto_original in parrafos_texto:
            parrafo_texto = parrafo_texto_original.strip()

            if parrafo_texto == '***': # Separador para la legalización
                documento.add_paragraph()
                documento.add_paragraph('************************************************************************************************************************')
                documento.add_paragraph()
                continue
            
            if not parrafo_texto:
                documento.add_paragraph()
                continue

            p = documento.add_paragraph()
            current_index = 0
            
            # Lógica para aplicar negritas: recorre el párrafo y añade runs (segmentos de texto)
            while current_index < len(parrafo_texto):
                found_match = False
                
                for bold_phrase in palabras_en_negrita_ordenadas:
                    
                    if parrafo_texto.find(bold_phrase, current_index) == current_index:
                        
                        # 1. Añadir la frase en negrita
                        run = p.add_run(bold_phrase)
                        run.bold = True
                        
                        current_index += len(bold_phrase)
                        found_match = True
                        break 
                
                if not found_match:
                    # 2. Si no es negrita, añadir el siguiente caracter sin formato
                    run = p.add_run(parrafo_texto[current_index])
                    run.bold = False
                    current_index += 1

        # 4. Guardar el documento
        documento.save(filename)
        messagebox.showinfo("Éxito", f"🎉 ¡El contrato ha sido generado y guardado como:\n{filename}!")

    except Exception as e:
        messagebox.showerror("Error de Generación", f"Ocurrió un error al generar el documento: {e}")

# --------------------------------------------------------------------------------------
# 4. CLASE DE LA INTERFAZ GRÁFICA (GUI con Tkinter)
# --------------------------------------------------------------------------------------

class GeneradorContratosApp:
    def __init__(self, master):
        self.master = master
        master.title("✍️ Generador Automático de Contratos DOCX")
        
        self.entry_vars = {} # Almacena las variables de la GUI
        
        # Crear un contenedor con scroll
        canvas = tk.Canvas(master)
        scrollbar = ttk.Scrollbar(master, orient="vertical", command=canvas.yview)
        self.scrollable_frame = ttk.Frame(canvas)

        self.scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(
                scrollregion=canvas.bbox("all")
            )
        )

        canvas.create_window((0, 0), window=self.scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)

        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        
        self.create_widgets()

    def create_widgets(self):
        row_index = 0
        current_section = None
        
        # Iterar a través de los campos definidos
        for marcador, label_text, section in CAMPOS:
            if section != current_section:
                # Crear un separador para la sección nueva
                ttk.Label(self.scrollable_frame, text=f"--- {section} ---", font=("Arial", 10, "bold"), foreground="blue").grid(row=row_index, column=0, columnspan=2, pady=(10, 2), sticky="w")
                row_index += 1
                current_section = section

            # Etiqueta (Label)
            ttk.Label(self.scrollable_frame, text=f"{label_text}:").grid(row=row_index, column=0, padx=5, pady=2, sticky="w")
            
            # Campo de entrada (Entry)
            var = tk.StringVar()
            entry = ttk.Entry(self.scrollable_frame, textvariable=var, width=50)
            entry.grid(row=row_index, column=1, padx=5, pady=2, sticky="ew")
            
            self.entry_vars[marcador] = var
            row_index += 1

        # Botón para generar el contrato
        ttk.Button(self.scrollable_frame, text="GENERAR CONTRATO DOCX", command=self.generar).grid(row=row_index + 1, column=0, columnspan=2, pady=20, padx=5)

    def generar(self):
        datos = {}
        for marcador, var in self.entry_vars.items():
            value = var.get().strip()
            
            # Validar que no haya campos vacíos
            if not value:
                messagebox.showwarning("Campos Incompletos", "Por favor, rellena todos los campos antes de generar el contrato.")
                return
            
            # Aplicar formato de capitalización/mayúsculas si corresponde
            if "LETRA" in marcador or "MAYÚS" in marcador or "ARRENDANTE" in marcador or "ARRENDATARIO" in marcador:
                datos[marcador] = value.upper()
            elif "Mes" in marcador or "mes" in marcador:
                 datos[marcador] = value.capitalize()
            else:
                datos[marcador] = value


        # Generar nombre de archivo (más limpio)
        nombre_arrendatario = datos['NOMBRE_ARRENDATARIO_COMPLETO'].replace(' ', '_').replace('.', '').replace(',', '')
        anio = datos.get('ANIO_INICIO', 'SIN_ANIO')
        filename = f"Contrato_{nombre_arrendatario}_{anio}.docx"

        generar_contrato_docx(datos, filename)
        
        # Opcional: limpiar los campos después de la generación
        # for var in self.entry_vars.values():
        #     var.set("")


if __name__ == "__main__":
    root = tk.Tk()
    app = GeneradorContratosApp(root)
    root.mainloop()